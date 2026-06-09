#!/usr/bin/env python3
"""Generate DOCX from the canonical case spec."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Microsoft JhengHei"
TITLE_COLOR = RGBColor(25, 43, 79)
HEADING_COLOR = RGBColor(49, 91, 163)
TEXT_COLOR = RGBColor(40, 40, 40)
MUTED_COLOR = RGBColor(95, 103, 115)


def load_spec(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_COLOR


def style_run(
    run, size: int, *, bold: bool = False, color: RGBColor | None = None
) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(
    doc: Document, text: str, size: int, color: RGBColor, *, center: bool = False
) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, size, bold=True, color=color)


def add_paragraph(
    doc: Document, text: str, *, size: int = 11, color: RGBColor | None = None
) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size, color=color or TEXT_COLOR)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        style_run(r, 11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(header)
        style_run(r, 10, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, "315BA3")
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            style_run(r, 10)
            set_cell_shading(cell, "F6F8FB" if row_idx % 2 == 0 else "FFFFFF")


def build_document(spec: dict[str, Any]) -> Document:
    doc = Document()
    set_defaults(doc)
    case = spec["case"]
    dp = spec["decision_pivot"]
    ctx = spec["company_industry_context"]
    dilemma = spec["core_dilemma"]
    evidence = spec["evidence"]
    appendix = spec.get("appendix", {})
    recommendation = spec.get("recommendation", {})

    add_heading(doc, case["title"], 20, TITLE_COLOR, center=True)
    if case.get("subtitle"):
        add_heading(doc, case["subtitle"], 11, MUTED_COLOR, center=True)
    if case.get("metadata"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, item in enumerate(case["metadata"]):
            if idx:
                sep = p.add_run(" | ")
                style_run(sep, 9, color=MUTED_COLOR)
            r = p.add_run(item)
            style_run(r, 9, color=MUTED_COLOR)

    add_heading(doc, "一、導言與決策切入點", 15, HEADING_COLOR)
    add_bullets(
        doc,
        [
            f"決策主角：{dp['decision_owner']}",
            f"核心決策：{dp['decision_question']}",
            f"緊迫性：{dp['urgency']}",
            f"延遲代價：{dp['delay_cost']}",
        ],
    )
    add_paragraph(doc, "時間軸：")
    add_bullets(doc, dp.get("milestones", []))

    add_heading(doc, "二、公司與產業總體檢", 15, HEADING_COLOR)
    add_paragraph(doc, "企業背景：")
    add_bullets(doc, ctx.get("company_background", []))
    add_paragraph(doc, "產業背景：")
    add_bullets(doc, ctx.get("industry_background", []))
    structural = ctx.get("structural_breakdown", {})
    for label, key in [
        ("市場與通路", "market_channel"),
        ("成本與供應鏈", "cost_supply_chain"),
        ("技術與組織", "technology_organization"),
    ]:
        if structural.get(key):
            add_paragraph(doc, f"{label}：")
            add_bullets(doc, structural[key])

    add_heading(doc, "三、核心衝突與根本問題診斷", 15, HEADING_COLOR)
    add_bullets(
        doc,
        [
            f"表層問題：{dilemma['surface_problem']}",
            f"根本問題：{dilemma['root_problem']}",
            f"核心張力：{dilemma['key_tension']}",
        ],
    )
    add_paragraph(doc, "5 Whys：")
    add_bullets(doc, dilemma.get("five_whys", []))
    add_paragraph(doc, "Trade-off：")
    add_bullets(doc, dilemma.get("trade_offs", []))

    add_heading(doc, "四、數據驗證與實例支持", 15, HEADING_COLOR)
    add_paragraph(doc, "關鍵數據：")
    add_bullets(doc, evidence.get("quantitative_signals", []))
    add_paragraph(doc, "內部檢查：")
    add_bullets(doc, evidence.get("internal_checks", []))
    add_paragraph(doc, "外部交叉檢查：")
    add_bullets(doc, evidence.get("external_checks", []))
    if evidence.get("open_issues"):
        add_paragraph(doc, "仍有疑點：")
        add_bullets(doc, evidence["open_issues"])

    add_heading(doc, "五、課堂思辨與行動決策", 15, HEADING_COLOR)
    for option in spec.get("options", []):
        add_paragraph(doc, option["title"], size=12, color=HEADING_COLOR)
        add_bullets(
            doc,
            [
                f"How：{option['how']}",
                f"Why：{option['why']}",
                f"Trade-off：{option['trade_off']}",
            ],
        )
    if recommendation:
        add_paragraph(doc, "建議方案", size=12, color=HEADING_COLOR)
        add_bullets(
            doc,
            [
                f"建議：{recommendation.get('recommended_option', '')}",
                f"理由：{recommendation.get('reason', '')}",
            ],
        )
        if recommendation.get("conditions"):
            add_paragraph(doc, "成立前提：")
            add_bullets(doc, recommendation["conditions"])

    add_heading(doc, "附錄與參考資料", 15, HEADING_COLOR)
    if appendix.get("references"):
        add_paragraph(doc, "參考來源：")
        add_bullets(doc, appendix["references"])
    if appendix.get("assumptions"):
        add_paragraph(doc, "關鍵假設：")
        add_bullets(doc, appendix["assumptions"])
    if appendix.get("discussion_questions"):
        add_paragraph(doc, "課堂討論題：")
        add_bullets(doc, appendix["discussion_questions"])
    if case.get("source_notes"):
        add_paragraph(doc, "Source Notes：")
        add_bullets(doc, case["source_notes"])

    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate DOCX from the canonical case spec."
    )
    parser.add_argument("spec", type=Path, help="Path to canonical case spec JSON.")
    parser.add_argument("output", type=Path, help="Path to output .docx file.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    spec = load_spec(args.spec)
    doc = build_document(spec)
    output_path = args.output.expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
